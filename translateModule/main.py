import googletrans
from pathlib import Path
import docx
from shutil import copy2
import sys


filePath = sys.argv[1]
fileSave = sys.argv[2]

lng_lst = []
for it in range(3, len(sys.argv)):
    lng_lst.append(sys.argv[it])


#############################################################################


def get_text_list_txt(path):

    text = []

    with open(path, 'r', encoding="utf-8") as file:
        for i in file:
            text.append(i)

    return tuple(text)


def get_text_list_docx(path):

    text = []
    doc = docx.Document(path)

    for i in doc.paragraphs:
        text.append(i.text)

    return tuple(text)


def translate_list(lt, index=0):

    translator = googletrans.Translator()
    text = []

    for i in lt:
        tr_res = ''
        try:
            tr_res = translator.translate(text=i, src=lng_lst[index], dest=lng_lst[index + 1]).text
        except Exception as exc:
            print("Warning:\t", exc)

        text.append(tr_res)

    if index + 1 == len(lng_lst) - 1:
        return tuple(text)
    else:
        return tuple(translate_list(text, index + 1))


def get_new_path():

    name = Path(filePath).name
    prefix = "U-file " if len(lng_lst) > 2 else "T-file "

    count = 1
    full_path = fileSave + '\\' + prefix + name

    print("WHILE")
    while Path(full_path).exists():
        print("IT")
        full_path = fileSave + '\\' + prefix + str(count) + name
        count += 1
    print("END")

    return full_path


def write_txt(lt):

    f_name = get_new_path()

    try:
        open(f_name, 'x')
    except Exception as exc:
        print("File already created!:\t", exc)

    with open(f_name, 'w', encoding='utf-8') as file:
        print(len(lt))
        for i in lt:
            file.write(i + '\n')


def write_docx(lt):

    f_name = get_new_path()
    copy2(filePath, f_name)

    doc = docx.Document(f_name)
    para = doc.paragraphs

    for i in range(0, len(para)):
        if not lt[i] == '':
            para[i].text = lt[i]

    doc.save(f_name)


#############################################################################


if __name__ == '__main__':

    # print(googletrans.LANGUAGES)

    ext = Path(filePath).suffix

    if ext == '.txt':
        write_txt(translate_list(get_text_list_txt(filePath)))
    elif ext == '.docx':
        write_docx(translate_list(get_text_list_docx(filePath)))
